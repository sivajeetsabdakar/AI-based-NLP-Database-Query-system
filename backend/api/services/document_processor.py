"""
Multi-Format Document Processing Pipeline
Handles PDF, DOCX, TXT, and CSV files with intelligent processing
"""
import logging
import os
import mimetypes
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import hashlib
import json
from pathlib import Path

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
import chardet

# Text processing
import re
from sentence_transformers import SentenceTransformer

# Database and storage
from .chromadb_service import get_chromadb_service
from .database_manager import get_database_manager
from .database_utils import get_database_utils
from .mistral_client import get_mistral_client

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Multi-format document processing with intelligent chunking"""
    
    def __init__(self):
        self.logger = logger
        self.chromadb_service = None
        self.db_manager = None
        self.db_utils = None
        self.embedding_model = None
        self.mistral_client = None
        
        # Document type patterns
        self.document_patterns = {
            "resume": [
                r"resume", r"cv", r"curriculum", r"experience", r"skills", 
                r"education", r"qualification", r"objective"
            ],
            "contract": [
                r"contract", r"agreement", r"terms", r"conditions", r"clause",
                r"legal", r"binding", r"signature", r"effective"
            ],
            "review": [
                r"review", r"performance", r"evaluation", r"assessment", 
                r"rating", r"feedback", r"appraisal", r"goals"
            ],
            "policy": [
                r"policy", r"procedure", r"guideline", r"rule", r"regulation",
                r"standard", r"protocol", r"compliance"
            ]
        }
        
        # Chunking strategies by document type
        self.chunking_strategies = {
            "resume": {
                "max_chunk_size": 512,
                "overlap": 50,
                "preserve_sections": True,
                "sections": ["skills", "experience", "education", "objective"]
            },
            "contract": {
                "max_chunk_size": 1024,
                "overlap": 100,
                "preserve_sections": True,
                "sections": ["clause", "section", "article", "paragraph"]
            },
            "review": {
                "max_chunk_size": 768,
                "overlap": 75,
                "preserve_sections": True,
                "sections": ["goal", "achievement", "feedback", "rating"]
            },
            "policy": {
                "max_chunk_size": 1024,
                "overlap": 100,
                "preserve_sections": True,
                "sections": ["section", "subsection", "clause", "article"]
            },
            "default": {
                "max_chunk_size": 512,
                "overlap": 50,
                "preserve_sections": False,
                "sections": []
            }
        }
    
    def initialize(self):
        """Initialize the document processor"""
        try:
            # Initialize services
            self.chromadb_service = get_chromadb_service()
            self.db_manager = get_database_manager()
            self.db_utils = get_database_utils()
            self.mistral_client = get_mistral_client()
            
            # Initialize embedding model
            self.logger.info("Loading sentence-transformers model...")
            try:
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                self.logger.info("Embedding model loaded successfully")
            except Exception as e:
                self.logger.error(f"Failed to load embedding model: {str(e)}")
                # Try alternative model
                try:
                    self.logger.info("Trying alternative embedding model...")
                    self.embedding_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
                    self.logger.info("Alternative embedding model loaded successfully")
                except Exception as e2:
                    self.logger.error(f"Failed to load alternative embedding model: {str(e2)}")
                    raise e2
            
            self.logger.info("Document processor initialized successfully")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize document processor: {str(e)}")
            raise
    
    def process_documents(self, file_paths: List[str], user_session_id: Optional[str] = None) -> Dict[str, Any]:
        """
        Process multiple documents with intelligent chunking
        
        Args:
            file_paths: List of file paths to process
            user_session_id: Optional user session ID
            
        Returns:
            Processing results with document IDs and status
        """
        try:
            self.logger.info(f"Processing {len(file_paths)} documents")
            
            processing_results = {
                "total_documents": len(file_paths),
                "processed_documents": 0,
                "failed_documents": 0,
                "document_ids": [],
                "errors": [],
                "processing_time": 0,
                "start_time": datetime.utcnow().isoformat()
            }
            
            start_time = datetime.utcnow()
            
            for file_path in file_paths:
                try:
                    # Process individual document
                    result = self.process_single_document(file_path, user_session_id)
                    
                    if result["success"]:
                        processing_results["processed_documents"] += 1
                        processing_results["document_ids"].append(result["document_id"])
                    else:
                        processing_results["failed_documents"] += 1
                        processing_results["errors"].append({
                            "file_path": file_path,
                            "error": result.get("error", "Unknown error")
                        })
                        
                except Exception as e:
                    self.logger.error(f"Failed to process document {file_path}: {str(e)}")
                    processing_results["failed_documents"] += 1
                    processing_results["errors"].append({
                        "file_path": file_path,
                        "error": str(e)
                    })
            
            processing_results["processing_time"] = (datetime.utcnow() - start_time).total_seconds()
            processing_results["end_time"] = datetime.utcnow().isoformat()
            
            self.logger.info(f"Document processing completed: {processing_results['processed_documents']} successful, {processing_results['failed_documents']} failed")
            
            return processing_results
            
        except Exception as e:
            self.logger.error(f"Document processing failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "total_documents": len(file_paths),
                "processed_documents": 0,
                "failed_documents": len(file_paths)
            }
    
    def process_single_document(self, file_path: str, user_session_id: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a single document with intelligent analysis
        
        Args:
            file_path: Path to the document file
            user_session_id: Optional user session ID
            
        Returns:
            Processing result with document ID and status
        """
        try:
            self.logger.info(f"Processing document: {file_path}")
            
            # Validate file
            if not os.path.exists(file_path):
                return {"success": False, "error": "File not found"}
            
            # Ensure services are initialized - try lazy initialization
            if not self.chromadb_service:
                self.logger.warning("ChromaDB service not initialized, attempting lazy initialization...")
                try:
                    self.chromadb_service = get_chromadb_service()
                except Exception as e:
                    self.logger.error(f"Failed to initialize ChromaDB service: {str(e)}")
                    return {"success": False, "error": "ChromaDB service not initialized"}
            
            if not self.db_manager:
                self.logger.warning("Database manager not initialized, attempting lazy initialization...")
                try:
                    self.db_manager = get_database_manager()
                except Exception as e:
                    self.logger.error(f"Failed to initialize database manager: {str(e)}")
                    return {"success": False, "error": "Database manager not initialized"}
            
            if not self.db_utils:
                self.logger.warning("Database utils not initialized, attempting lazy initialization...")
                try:
                    self.db_utils = get_database_utils()
                except Exception as e:
                    self.logger.error(f"Failed to initialize database utils: {str(e)}")
                    return {"success": False, "error": "Database utils not initialized"}
            
            if not self.mistral_client:
                self.logger.warning("Mistral client not initialized, attempting lazy initialization...")
                try:
                    self.mistral_client = get_mistral_client()
                except Exception as e:
                    self.logger.warning(f"Failed to initialize Mistral client: {str(e)}")
                    # Mistral client is optional, continue without it
            
            if not self.embedding_model:
                self.logger.warning("Embedding model not initialized, attempting lazy initialization...")
                try:
                    self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                    self.logger.info("Embedding model loaded successfully (lazy initialization)")
                except Exception as e:
                    self.logger.error(f"Failed to load embedding model: {str(e)}")
                    # Try alternative model
                    try:
                        self.logger.info("Trying alternative embedding model...")
                        self.embedding_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
                        self.logger.info("Alternative embedding model loaded successfully (lazy initialization)")
                    except Exception as e2:
                        self.logger.error(f"Failed to load alternative embedding model: {str(e2)}")
                        return {"success": False, "error": "Embedding model not initialized"}
            
            # Detect document type
            doc_type = self._detect_document_type(file_path)
            
            # Extract content
            content = self._extract_content(file_path, doc_type)
            if not content:
                return {"success": False, "error": "Failed to extract content"}
            
            # Analyze document structure
            structure = self._analyze_document_structure(content, doc_type)
            
            # Generate intelligent chunks
            chunks = self._generate_intelligent_chunks(content, doc_type, structure)
            
            if not chunks:
                return {"success": False, "error": "Failed to generate document chunks"}
            
            # Generate embeddings
            embeddings = self._generate_embeddings(chunks)
            
            if not embeddings:
                return {"success": False, "error": "Failed to generate embeddings"}
            
            # Store in ChromaDB
            document_id = self._store_document_chunks(file_path, doc_type, chunks, embeddings, user_session_id)
            
            # Store metadata in database
            self._store_document_metadata(file_path, doc_type, document_id, user_session_id)
            
            self.logger.info(f"Document processed successfully: {document_id}")
            
            return {
                "success": True,
                "document_id": document_id,
                "doc_type": doc_type,
                "chunks_count": len(chunks),
                "structure": structure
            }
            
        except Exception as e:
            self.logger.error(f"Failed to process document {file_path}: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _detect_document_type(self, file_path: str) -> str:
        """Detect document type based on file extension and content"""
        try:
            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            
            # Map extensions to document types
            extension_map = {
                '.pdf': 'pdf',
                '.docx': 'docx',
                '.doc': 'docx',
                '.txt': 'txt',
                '.csv': 'csv'
            }
            
            file_type = extension_map.get(file_ext, 'unknown')
            
            # For text-based files, analyze content to determine document purpose
            if file_type in ['txt', 'pdf', 'docx']:
                # Extract sample content for analysis
                sample_content = self._extract_sample_content(file_path, file_type)
                if sample_content:
                    # Analyze content to determine document purpose
                    purpose = self._analyze_document_purpose(sample_content)
                    return f"{file_type}_{purpose}"
            
            return file_type
            
        except Exception as e:
            self.logger.error(f"Failed to detect document type: {str(e)}")
            return "unknown"
    
    def _extract_sample_content(self, file_path: str, file_type: str) -> str:
        """Extract sample content for document type analysis"""
        try:
            if file_type == 'pdf':
                with pdfplumber.open(file_path) as pdf:
                    # Extract first page content
                    if pdf.pages:
                        page = pdf.pages[0]
                        return page.extract_text()[:1000]  # First 1000 characters
            
            elif file_type == 'docx':
                doc = DocxDocument(file_path)
                # Extract first paragraph
                if doc.paragraphs:
                    return doc.paragraphs[0].text[:1000]
            
            elif file_type == 'txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(1000)
            
            return ""
            
        except Exception as e:
            self.logger.error(f"Failed to extract sample content: {str(e)}")
            return ""
    
    def _analyze_document_purpose(self, content: str) -> str:
        """Analyze document content to determine purpose"""
        try:
            content_lower = content.lower()
            
            # Check for document purpose patterns
            for purpose, patterns in self.document_patterns.items():
                matches = sum(1 for pattern in patterns if re.search(pattern, content_lower))
                if matches >= 2:  # Threshold for purpose detection
                    return purpose
            
            return "general"
            
        except Exception as e:
            self.logger.error(f"Failed to analyze document purpose: {str(e)}")
            return "general"
    
    def _extract_content(self, file_path: str, doc_type: str) -> str:
        """Extract content from document based on type"""
        try:
            if doc_type.startswith('pdf'):
                return self._extract_pdf_content(file_path)
            elif doc_type.startswith('docx'):
                return self._extract_docx_content(file_path)
            elif doc_type.startswith('txt'):
                return self._extract_txt_content(file_path)
            elif doc_type.startswith('csv'):
                return self._extract_csv_content(file_path)
            else:
                return self._extract_generic_content(file_path)
                
        except Exception as e:
            self.logger.error(f"Failed to extract content from {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract content from PDF file using Mistral OCR"""
        try:
            # First try Mistral OCR for better text extraction
            if self.mistral_client:
                try:
                    self.logger.info(f"Using Mistral OCR for PDF extraction: {file_path}")
                    ocr_result = self.mistral_client.extract_text_from_pdf_ocr(file_path)
                    
                    if ocr_result["success"] and ocr_result["text"]:
                        self.logger.info(f"Mistral OCR extracted {len(ocr_result['text'])} characters from {ocr_result['pages_processed']} pages")
                        return ocr_result["text"]
                    else:
                        self.logger.warning(f"Mistral OCR failed: {ocr_result.get('error', 'Unknown error')}")
                        
                except Exception as e:
                    self.logger.warning(f"Mistral OCR failed, falling back to traditional methods: {str(e)}")
            else:
                self.logger.warning("Mistral client not available, using traditional PDF extraction")
            
            # Fallback to traditional PDF extraction methods
            content = ""
            
            # Try pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(file_path) as pdf:
                    if not pdf.pages:
                        self.logger.warning("PDF has no pages")
                        return ""
                    
                    for i, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                content += page_text + "\n"
                        except Exception as e:
                            self.logger.warning(f"Failed to extract text from page {i}: {str(e)}")
                            continue
                            
            except Exception as e:
                self.logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
                
                # Fallback to PyPDF2
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        if not pdf_reader.pages:
                            self.logger.warning("PDF has no pages (PyPDF2)")
                            return ""
                            
                        for i, page in enumerate(pdf_reader.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    content += page_text + "\n"
                            except Exception as e:
                                self.logger.warning(f"Failed to extract text from page {i} (PyPDF2): {str(e)}")
                                continue
                except Exception as e2:
                    self.logger.error(f"Both PDF extraction methods failed: {str(e2)}")
                    return ""
            
            if not content.strip():
                self.logger.warning("No content extracted from PDF")
                return ""
                
            return content.strip()
            
        except Exception as e:
            self.logger.error(f"Failed to extract PDF content: {str(e)}")
            return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extract content from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content += cell.text + " "
                    content += "\n"
            
            return content.strip()
            
        except Exception as e:
            self.logger.error(f"Failed to extract DOCX content: {str(e)}")
            return ""
    
    def _extract_txt_content(self, file_path: str) -> str:
        """Extract content from TXT file with encoding detection"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            self.logger.error(f"Failed to extract TXT content: {str(e)}")
            return ""
    
    def _extract_csv_content(self, file_path: str) -> str:
        """Extract content from CSV file"""
        try:
            # Read CSV with pandas
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = ""
            
            # Add column headers
            content += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"
            
            # Add sample data
            for index, row in df.head(10).iterrows():  # First 10 rows
                row_text = " | ".join([str(value) for value in row.values if pd.notna(value)])
                content += row_text + "\n"
            
            return content.strip()
            
        except Exception as e:
            self.logger.error(f"Failed to extract CSV content: {str(e)}")
            return ""
    
    def _extract_generic_content(self, file_path: str) -> str:
        """Extract content from unknown file type"""
        try:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            self.logger.error(f"Failed to extract generic content: {str(e)}")
            return ""
    
    def _analyze_document_structure(self, content: str, doc_type: str) -> Dict[str, Any]:
        """Analyze document structure for intelligent chunking"""
        try:
            structure = {
                "sections": [],
                "paragraphs": [],
                "tables": [],
                "lists": [],
                "headings": []
            }
            
            # Split content into paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            structure["paragraphs"] = paragraphs
            
            # Detect headings (lines that are shorter and might be titles)
            for i, para in enumerate(paragraphs):
                if len(para) < 100 and para.isupper() or para.endswith(':'):
                    structure["headings"].append({"text": para, "position": i})
            
            # Detect lists (lines starting with bullets or numbers)
            for i, para in enumerate(paragraphs):
                if re.match(r'^[\s]*[•\-\*\d+\.]', para):
                    structure["lists"].append({"text": para, "position": i})
            
            # Detect sections based on document type
            if doc_type.startswith('resume'):
                structure["sections"] = self._detect_resume_sections(paragraphs)
            elif doc_type.startswith('contract'):
                structure["sections"] = self._detect_contract_sections(paragraphs)
            elif doc_type.startswith('review'):
                structure["sections"] = self._detect_review_sections(paragraphs)
            elif doc_type.startswith('policy'):
                structure["sections"] = self._detect_policy_sections(paragraphs)
            
            return structure
            
        except Exception as e:
            self.logger.error(f"Failed to analyze document structure: {str(e)}")
            return {"sections": [], "paragraphs": [], "tables": [], "lists": [], "headings": []}
    
    def _detect_resume_sections(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """Detect resume sections"""
        sections = []
        section_keywords = {
            "objective": ["objective", "summary", "profile"],
            "experience": ["experience", "employment", "work history", "career"],
            "education": ["education", "academic", "qualification", "degree"],
            "skills": ["skills", "competencies", "expertise", "abilities"],
            "projects": ["projects", "portfolio", "work samples"]
        }
        
        for i, para in enumerate(paragraphs):
            para_lower = para.lower()
            for section_type, keywords in section_keywords.items():
                if any(keyword in para_lower for keyword in keywords):
                    sections.append({
                        "type": section_type,
                        "text": para,
                        "position": i
                    })
                    break
        
        return sections
    
    def _detect_contract_sections(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """Detect contract sections"""
        sections = []
        section_keywords = {
            "parties": ["parties", "agreement between", "contracting parties"],
            "terms": ["terms", "conditions", "provisions"],
            "payment": ["payment", "compensation", "fees", "salary"],
            "termination": ["termination", "expiration", "end of agreement"],
            "signature": ["signature", "execution", "signed"]
        }
        
        for i, para in enumerate(paragraphs):
            para_lower = para.lower()
            for section_type, keywords in section_keywords.items():
                if any(keyword in para_lower for keyword in keywords):
                    sections.append({
                        "type": section_type,
                        "text": para,
                        "position": i
                    })
                    break
        
        return sections
    
    def _detect_review_sections(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """Detect performance review sections"""
        sections = []
        section_keywords = {
            "goals": ["goals", "objectives", "targets"],
            "achievements": ["achievements", "accomplishments", "results"],
            "feedback": ["feedback", "comments", "observations"],
            "rating": ["rating", "score", "evaluation"],
            "development": ["development", "improvement", "growth"]
        }
        
        for i, para in enumerate(paragraphs):
            para_lower = para.lower()
            for section_type, keywords in section_keywords.items():
                if any(keyword in para_lower for keyword in keywords):
                    sections.append({
                        "type": section_type,
                        "text": para,
                        "position": i
                    })
                    break
        
        return sections
    
    def _detect_policy_sections(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """Detect policy document sections"""
        sections = []
        section_keywords = {
            "purpose": ["purpose", "objective", "scope"],
            "policy": ["policy", "procedure", "guideline"],
            "compliance": ["compliance", "violation", "enforcement"],
            "review": ["review", "update", "revision"],
            "contact": ["contact", "questions", "inquiries"]
        }
        
        for i, para in enumerate(paragraphs):
            para_lower = para.lower()
            for section_type, keywords in section_keywords.items():
                if any(keyword in para_lower for keyword in keywords):
                    sections.append({
                        "type": section_type,
                        "text": para,
                        "position": i
                    })
                    break
        
        return sections
    
    def _generate_intelligent_chunks(self, content: str, doc_type: str, structure: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate intelligent chunks based on document type and structure"""
        try:
            # Get chunking strategy for document type
            base_type = doc_type.split('_')[0] if '_' in doc_type else doc_type
            strategy = self.chunking_strategies.get(base_type, self.chunking_strategies["default"])
            
            chunks = []
            
            if strategy["preserve_sections"] and structure["sections"]:
                # Section-based chunking
                chunks = self._chunk_by_sections(content, structure["sections"], strategy)
            else:
                # Standard chunking
                chunks = self._chunk_by_size(content, strategy)
            
            # Add metadata to chunks
            for i, chunk in enumerate(chunks):
                chunk["chunk_id"] = f"{doc_type}_{i}"
                chunk["chunk_index"] = i
                chunk["doc_type"] = doc_type
                chunk["total_chunks"] = len(chunks)
                chunk["created_at"] = datetime.utcnow().isoformat()
            
            self.logger.info(f"Generated {len(chunks)} intelligent chunks for {doc_type}")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Failed to generate intelligent chunks: {str(e)}")
            return []
    
    def _chunk_by_sections(self, content: str, sections: List[Dict[str, Any]], strategy: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Chunk content by sections"""
        chunks = []
        
        for section in sections:
            section_text = section["text"]
            
            # If section is too large, split it further
            if len(section_text) > strategy["max_chunk_size"]:
                sub_chunks = self._split_large_text(section_text, strategy)
                for sub_chunk in sub_chunks:
                    chunks.append({
                        "text": sub_chunk,
                        "section_type": section["type"],
                        "section_position": section["position"],
                        "chunk_type": "section"
                    })
            else:
                chunks.append({
                    "text": section_text,
                    "section_type": section["type"],
                    "section_position": section["position"],
                    "chunk_type": "section"
                })
        
        return chunks
    
    def _chunk_by_size(self, content: str, strategy: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Chunk content by size with overlap"""
        chunks = []
        max_size = strategy["max_chunk_size"]
        overlap = strategy["overlap"]
        
        # Split content into sentences
        sentences = re.split(r'[.!?]+', content)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        current_chunk = ""
        chunk_sentences = []
        
        for sentence in sentences:
            # Check if adding this sentence would exceed max size
            if len(current_chunk) + len(sentence) > max_size and current_chunk:
                # Save current chunk
                chunks.append({
                    "text": current_chunk.strip(),
                    "chunk_type": "size_based",
                    "sentences": chunk_sentences.copy()
                })
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
                chunk_sentences = [sentence]
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                chunk_sentences.append(sentence)
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                "text": current_chunk.strip(),
                "chunk_type": "size_based",
                "sentences": chunk_sentences
            })
        
        return chunks
    
    def _split_large_text(self, text: str, strategy: Dict[str, Any]) -> List[str]:
        """Split large text into smaller chunks"""
        max_size = strategy["max_chunk_size"]
        overlap = strategy["overlap"]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Find last sentence boundary before end
            last_period = text.rfind('.', start, end)
            last_exclamation = text.rfind('!', start, end)
            last_question = text.rfind('?', start, end)
            
            boundary = max(last_period, last_exclamation, last_question)
            
            if boundary > start:
                end = boundary + 1
            
            chunks.append(text[start:end])
            start = end - overlap if end > overlap else end
        
        return chunks
    
    def _generate_embeddings(self, chunks: List[Dict[str, Any]]) -> List[List[float]]:
        """Generate embeddings for document chunks"""
        try:
            if not self.embedding_model:
                self.logger.error("Embedding model not initialized")
                raise ValueError("Embedding model not initialized")
            
            if not chunks:
                self.logger.warning("No chunks provided for embedding generation")
                return []
            
            # Extract text from chunks
            texts = [chunk["text"] for chunk in chunks if chunk.get("text")]
            
            if not texts:
                self.logger.warning("No valid text chunks found for embedding generation")
                return []
            
            # Generate embeddings in batches
            batch_size = 32
            embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                if batch_texts:  # Ensure batch is not empty
                    batch_embeddings = self.embedding_model.encode(batch_texts)
                    embeddings.extend(batch_embeddings.tolist())
            
            self.logger.info(f"Generated {len(embeddings)} embeddings")
            return embeddings
            
        except Exception as e:
            self.logger.error(f"Failed to generate embeddings: {str(e)}")
            return []
    
    def _store_document_chunks(self, file_path: str, doc_type: str, chunks: List[Dict[str, Any]], embeddings: List[List[float]], user_session_id: Optional[str] = None) -> str:
        """Store document chunks in ChromaDB"""
        try:
            if not self.chromadb_service or not self.chromadb_service.client:
                self.logger.error("ChromaDB service not initialized")
                raise ValueError("ChromaDB service not initialized")
            
            if not chunks:
                self.logger.warning("No chunks provided for storage")
                raise ValueError("No chunks provided for storage")
            
            if not embeddings:
                self.logger.warning("No embeddings provided for storage")
                raise ValueError("No embeddings provided for storage")
            
            if len(chunks) != len(embeddings):
                self.logger.error(f"Mismatch between chunks ({len(chunks)}) and embeddings ({len(embeddings)})")
                raise ValueError(f"Mismatch between chunks ({len(chunks)}) and embeddings ({len(embeddings)})")
            
            # Generate document ID
            file_hash = hashlib.sha256(file_path.encode()).hexdigest()
            document_id = f"{doc_type}_{file_hash[:16]}"
            
            # Determine collection based on document type
            collection_name = self._get_collection_name(doc_type)
            
            # Prepare documents, embeddings, and metadata
            documents = []
            metadatas = []
            ids = []
            
            for i, chunk in enumerate(chunks):
                if not chunk.get("text"):
                    self.logger.warning(f"Skipping empty chunk {i}")
                    continue
                    
                documents.append(chunk["text"])
                
                metadata = {
                    "document_id": document_id,
                    "file_path": file_path,
                    "doc_type": doc_type,
                    "chunk_index": i,
                    "chunk_type": chunk.get("chunk_type", "unknown"),
                    "section_type": chunk.get("section_type", "unknown"),
                    "created_at": chunk.get("created_at", datetime.utcnow().isoformat()),
                    "user_session_id": user_session_id or "default"
                }
                
                # Remove any None values from metadata
                metadata = {k: v for k, v in metadata.items() if v is not None}
                
                # Add section-specific metadata
                if "section_position" in chunk:
                    metadata["section_position"] = chunk["section_position"]
                
                metadatas.append(metadata)
                ids.append(f"{document_id}_chunk_{i}")
            
            if not documents:
                self.logger.error("No valid documents to store")
                raise ValueError("No valid documents to store")
            
            # Ensure embeddings match documents
            if len(embeddings) > len(documents):
                embeddings = embeddings[:len(documents)]
            elif len(embeddings) < len(documents):
                self.logger.error(f"Not enough embeddings for documents: {len(embeddings)} vs {len(documents)}")
                raise ValueError(f"Not enough embeddings for documents: {len(embeddings)} vs {len(documents)}")
            
            # Store in ChromaDB
            collection = self.chromadb_service.client.get_or_create_collection(collection_name)
            collection.add(
                documents=documents,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            self.logger.info(f"Stored {len(documents)} chunks in ChromaDB collection '{collection_name}'")
            return document_id
            
        except Exception as e:
            self.logger.error(f"Failed to store document chunks: {str(e)}")
            raise
    
    def _get_collection_name(self, doc_type: str) -> str:
        """Get ChromaDB collection name based on document type"""
        if doc_type.startswith('resume'):
            return "resume_chunks"
        elif doc_type.startswith('contract'):
            return "contract_chunks"
        elif doc_type.startswith('review'):
            return "review_chunks"
        elif doc_type.startswith('policy'):
            return "policy_chunks"
        else:
            return "employee_documents"
    
    def _store_document_metadata(self, file_path: str, doc_type: str, document_id: str, user_session_id: Optional[str] = None):
        """Store document metadata in database"""
        try:
            if not self.db_manager:
                return
            
            with self.db_manager.get_session() as session:
                from ..models.database_models import DocumentMetadata
                
                # Get file information
                file_size = os.path.getsize(file_path)
                filename = os.path.basename(file_path)
                
                # Create document metadata record
                doc_metadata = DocumentMetadata(
                    filename=filename,
                    file_type=doc_type,
                    file_size=file_size,
                    processing_status="completed",
                    chroma_id=document_id,
                    user_session_id=user_session_id,
                    document_metadata={
                        "file_path": file_path,
                        "doc_type": doc_type,
                        "processing_time": datetime.utcnow().isoformat()
                    }
                )
                
                session.add(doc_metadata)
                session.commit()
                
                self.logger.info(f"Stored document metadata for {document_id}")
                
        except Exception as e:
            self.logger.error(f"Failed to store document metadata: {str(e)}")
    
    def get_processing_status(self, document_id: str) -> Dict[str, Any]:
        """Get processing status for a document"""
        try:
            if not self.db_manager:
                return {"status": "error", "error": "Database not initialized"}
            
            with self.db_manager.get_session() as session:
                from ..models.database_models import DocumentMetadata
                
                doc_metadata = session.query(DocumentMetadata).filter(
                    DocumentMetadata.chroma_id == document_id
                ).first()
                
                if not doc_metadata:
                    return {"status": "not_found"}
                
                return {
                    "status": "found",
                    "document_id": document_id,
                    "filename": doc_metadata.filename,
                    "file_type": doc_metadata.file_type,
                    "processing_status": doc_metadata.processing_status,
                    "created_at": doc_metadata.created_at.isoformat() if doc_metadata.created_at else None
                }
                
        except Exception as e:
            self.logger.error(f"Failed to get processing status: {str(e)}")
            return {"status": "error", "error": str(e)}
    
    def search_documents(self, query: str, doc_type: Optional[str] = None, limit: int = 10) -> List[Dict[str, Any]]:
        """Search documents using vector similarity"""
        try:
            if not self.chromadb_service or not self.chromadb_service.client:
                raise ValueError("ChromaDB service not initialized")
            
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0].tolist()
            
            # Determine collections to search
            collections_to_search = []
            if doc_type:
                collection_name = self._get_collection_name(doc_type)
                collections_to_search.append(collection_name)
            else:
                collections_to_search = [
                    "employee_documents",
                    "resume_chunks", 
                    "contract_chunks",
                    "review_chunks",
                    "policy_chunks"
                ]
            
            results = []
            
            for collection_name in collections_to_search:
                try:
                    collection = self.chromadb_service.client.get_collection(collection_name)
                    
                    # Search in collection
                    search_results = collection.query(
                        query_embeddings=[query_embedding],
                        n_results=limit
                    )
                    
                    # Process results
                    if search_results["documents"] and search_results["documents"][0]:
                        for i, doc in enumerate(search_results["documents"][0]):
                            result = {
                                "document": doc,
                                "metadata": search_results["metadatas"][0][i] if search_results["metadatas"] else {},
                                "distance": search_results["distances"][0][i] if search_results["distances"] else 0.0,
                                "collection": collection_name
                            }
                            results.append(result)
                            
                except Exception as e:
                    self.logger.warning(f"Failed to search collection {collection_name}: {str(e)}")
                    continue
            
            # Sort by distance (lower is better)
            results.sort(key=lambda x: x["distance"])
            
            return results[:limit]
            
        except Exception as e:
            self.logger.error(f"Failed to search documents: {str(e)}")
            return []

# Global document processor instance
document_processor: Optional[DocumentProcessor] = None

def get_document_processor() -> DocumentProcessor:
    """Get the global document processor instance"""
    if document_processor is None:
        raise RuntimeError("Document processor not initialized")
    return document_processor

def initialize_document_processor() -> DocumentProcessor:
    """Initialize the global document processor"""
    global document_processor
    document_processor = DocumentProcessor()
    document_processor.initialize()
    return document_processor

def get_or_initialize_document_processor() -> DocumentProcessor:
    """Get or initialize the global document processor with lazy initialization"""
    global document_processor
    if document_processor is None:
        logger.info("Document processor not initialized, creating new instance...")
        document_processor = DocumentProcessor()
        try:
            document_processor.initialize()
            logger.info("Document processor initialized successfully")
        except Exception as e:
            logger.warning(f"Document processor initialization failed: {str(e)}")
            logger.info("Document processor will use lazy initialization")
    return document_processor